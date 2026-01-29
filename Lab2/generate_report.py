from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_front_page(document):
    # University of Ghana header
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('UNIVERSITY OF GHANA')
    run.bold = True
    run.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DEPARTMENT OF COMPUTER SCIENCE\nSCHOOL OF PHYSICAL AND MATHEMATICAL SCIENCE')
    run.font.size = Pt(12)

    document.add_paragraph('\n' * 5)

    # Course Title
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DCIT 403 – Designing Intelligent Agent')
    run.bold = True
    run.font.size = Pt(14)

    # Lab Title
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LAB 2: PERCEPTION AND ENVIRONMENT MODELING')
    run.bold = True
    run.font.size = Pt(18)

    document.add_paragraph('\n' * 10)

    # Student Info
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Name: ____________________\nID: ____________________\nDate: ' + os.popen('date /t').read().strip())
    run.font.size = Pt(11)

    document.add_page_break()

def create_report():
    document = Document()
    add_front_page(document)

    document.add_heading('1. Objective', level=1)
    document.add_paragraph(
        'The objective of this laboratory is to implement agent perception within a simulated disaster environment. '
        'This involves modeling environmental events and enabling a SensorAgent to detect and log these changes.'
    )

    document.add_heading('2. SensorAgent Implementation', level=1)
    document.add_paragraph(
        'The SensorAgent uses a PeriodicBehaviour to poll the DisasterEnvironment. It is designed to operate '
        'asynchronously and handle data parsing from the environment state.'
    )
    document.add_paragraph('[Source Files: Lab2/sensor_agent.py, Lab2/environment.py]')

    document.add_heading('3. Explanation of Percepts', level=1)
    document.add_paragraph(
        'Perception is modeled as a set of structured data representing environmental anomalies:'
    )
    document.add_paragraph('• Type (fire, flood, earthquake, etc.): Categorical perception of the disaster class.', style='List Bullet')
    document.add_paragraph('• Severity (low, medium, high, critical): Evaluation of the impact magnitude.', style='List Bullet')
    document.add_paragraph('• Location (x, y): Spatial coordinates of the event within the simulated 100x100 grid.', style='List Bullet')
    document.add_paragraph('• Timestamp: Temporal data indicating the exact time of detection.', style='List Bullet')

    document.add_heading('4. Event Logs (Execution Trace)', level=1)
    document.add_paragraph('The following represents the event logs generated during the simulation session:')
    
    log_path = os.path.join('event_log.txt')
    try:
        if not os.path.exists(log_path):
             # Fallback for folder structure
             log_path = os.path.join('Lab2', 'event_log.txt')
             
        with open(log_path, 'r') as f:
            logs = f.read()
            p = document.add_paragraph(logs)
            p.style = 'No Spacing' 
    except FileNotFoundError:
        document.add_paragraph("Simulated Trace: [Logs missing - Run simulate.py to generate]")

    document.add_heading('5. Troubleshooting & Debugging', level=1)
    document.add_paragraph(
        'During development, pathing issues were encountered when importing environment.py into sensor_agent.py. '
        'Resolved by updating sys.path or using relative package imports.'
    )

    document.save('Lab2_Report.docx')
    print("Professional Report generated: Lab2_Report.docx")

if __name__ == "__main__":
    create_report()
