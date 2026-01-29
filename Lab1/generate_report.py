from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_front_page(document):
    # University of Ghana header
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('UNIVERSITY OF GHANA')
    run.bold = True
    run.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DEPARTMENT OF COMPUTER SCIENCE\nSCHOOL OF PHYSICAL AND MATHEMATICAL SCIENCE')
    run.font.size = Pt(12)

    document.add_paragraph('\n' * 5)

    # Course Title
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DCIT 403 – Designing Intelligent Agent')
    run.bold = True
    run.font.size = Pt(14)

    # Lab Title
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LAB 1: ENVIRONMENT AND AGENT PLATFORM SETUP')
    run.bold = True
    run.font.size = Pt(18)

    document.add_paragraph('\n' * 10)

    # Student Info
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Name: ____________________\nID: ____________________\nDate: ' + os.popen('date /t').read().strip())
    run.font.size = Pt(11)

    document.add_page_break()

def create_report():
    document = Document()
    add_front_page(document)

    document.add_heading('1. Introduction', level=1)
    document.add_paragraph(
        'This laboratory session focuses on the configuration of the Python agent development environment '
        'using the Smart Python Agent Development Environment (SPADE) framework and XMPP protocol.'
    )

    document.add_heading('2. Environment Setup Report', level=1)
    document.add_paragraph(
        'The setup was conducted using a local development environment. Detailed steps taken include:'
    )
    document.add_paragraph('• Python 3.9+ Installation & Verification.', style='List Bullet')
    document.add_paragraph('• SPADE Framework Installation via pip.', style='List Bullet')
    document.add_paragraph('• Deployment of Prosody XMPP Server via Docker containerization.', style='List Bullet')
    document.add_paragraph('• Configuration of agent JID and credentials for local authentication.', style='List Bullet')

    document.add_heading('3. Source Code (basic_agent.py)', level=1)
    document.add_paragraph('The agent was implemented as an asynchronous class inheriting from spade.agent.Agent, '
                           'utilizing a CyclicBehaviour for periodic status updates.')
    
    # In a real scenario, we might paste the code here or just reference it.
    document.add_paragraph('[Refer to Lab1/basic_agent.py in the repository for full source code]')

    document.add_heading('4. Setup Challenges & Troubleshooting', level=1)
    
    document.add_heading('4.1 Python Version Compatibility', level=2)
    document.add_paragraph(
        'Observation: On systems running Python 3.14+, pre-built wheels for dependencies like lxml and aiohttp '
        'may unavailable. This results in pip attempting to build from source, which requires Microsoft Visual C++ Build Tools.'
    )
    document.add_paragraph(
        'Solution: Use a stable Python version (3.9-3.11) or ensure C++ build tools are installed. For this lab, '
        'GitHub Codespaces is the recommended platform as it provides a pre-configured environment.'
    )

    document.add_heading('4.2 XMPP Connection Errors', level=2)
    document.add_paragraph(
        'Observation: Agents may fail to connect if the XMPP server is not reachable or if registration is disabled.'
    )
    document.add_paragraph(
        'Solution: Ensure the Docker container is running (docker-compose up -d) and utilize "auto_register=True" '
        'in the agent.start() method to dynamically create accounts on the local server.'
    )

    document.add_heading('5. Conclusion', level=1)
    document.add_paragraph(
        'The environment setup is complete. The agent platform is operational and capable of supporting '
        'autonomous multi-agent communication for the subsequent laboratory tasks.'
    )

    document.save('Lab1_Report.docx')
    print("Professional Report generated: Lab1_Report.docx")

if __name__ == "__main__":
    create_report()
