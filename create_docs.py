import os
from docx import Document
from docx.shared import Pt

def create_lab_doc(lab_name, description, student_id, repo_link, log_files, output_filename, log_dir):
    doc = Document()
    doc.add_heading(f'{lab_name} Submission', 0)
    
    # Student Info
    p_info = doc.add_paragraph()
    p_info.add_run(f"Student ID: ").bold = True
    p_info.add_run(f"{student_id}\n")
    p_info.add_run(f"GitHub Repository: ").bold = True
    p_info.add_run(f"{repo_link}\n")
    
    # Description
    doc.add_heading('Description of Work Done', level=1)
    doc.add_paragraph(description)
        
    # Logs
    doc.add_heading('Execution Logs (Truncated)', level=1)
    for log_file in log_files:
        log_path = os.path.join(log_dir, log_file)
        if os.path.exists(log_path):
            doc.add_heading(log_file, level=2)
            with open(log_path, 'r', encoding='utf-8') as file:
                lines = []
                for i, line in enumerate(file):
                    if i >= 500: # Reduced to 500 lines for conciseness
                        lines.append("... [LOG TRUNCATED FOR SUBMISSION] ...\n")
                        break
                    lines.append(line)
                content = "".join(lines)
            
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)

    doc.save(output_filename)
    print(f"Saved {output_filename}")

if __name__ == "__main__":
    STUDENT_ID = "11041617"
    REPO_LINK = "https://github.com/Prudentkurler/DCIT403-Labs.git"
    
    # Lab 3
    lab3_desc = (
        "In Lab 3, I implemented an agentic simulation of a medical emergency response system "
        "using python's `transitions` library for Finite State Machines (FSMs). "
        "I created a 'Medical Agent' to monitor vital signs and a 'Rescue Agent' to respond to emergencies. "
        "The agents communicate via an event-driven publish/subscribe system. "
        "The system transitions between Normal, Alert, and Critical states based on monitored health metrics, "
        "triggering appropriate rescue protocols."
    )
    create_lab_doc(
        'Lab 3', 
        lab3_desc,
        STUDENT_ID,
        REPO_LINK,
        ['lab3_execution.log', 'lab3_execution_trace.txt'], 
        'Lab3_Submission.docx',
        'lab3'
    )
    
    # Lab 4
    lab4_desc = (
        "In Lab 4, I upgraded the emergency response system to utilize the SPADE framework "
        "(Smart Python Agent Development Environment) for proper multi-agent communication. "
        "I replaced the custom FSM simulation with proper SPADE Agents (Coordinator, Field, and Communicating Agents) "
        "running in an asynchronous XMPP environment (via a local Prosody server). "
        "The agents now communicate using standardized FIPA-ACL messages (Inform, Request, Agree, Failure, etc.) "
        "and handle complex coordination workflows asynchronously."
    )
    create_lab_doc(
        'Lab 4', 
        lab4_desc,
        STUDENT_ID,
        REPO_LINK,
        ['lab4_execution.log', 'lab4_message_log.json', 'lab4_message_log.txt'], 
        'Lab4_Submission.docx',
        'lab4'
    )
